from docx import Document
import json

def extract_qa_from_docx(file_path):
    doc = Document(file_path)
    data = []
    current_question = None
    answer = ""  # ✅ Initialize answer before use

    for para in doc.paragraphs:
        text = para.text.strip()
        if text.endswith("?"):  # Identify questions
            if current_question:
                data.append({"question": current_question, "answer": answer.strip()})
            current_question = text
            answer = ""  # ✅ Reset answer when a new question is found
        elif text:
            answer += " " + text  # ✅ Append answer text properly

    if current_question:
        data.append({"question": current_question, "answer": answer.strip()})

    return data
# Function to extract quotes from the quotes file
def extract_quotes_from_docx(file_path):
    doc = Document(file_path)
    quotes = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    return quotes
# ✅ Call the function with your file path
qa_data = extract_qa_from_docx(r"C:\Users\ESTHER KARANJA\Desktop\ChamaBot\chama resources.docx")
quotes = extract_quotes_from_docx(r"C:\Users\ESTHER KARANJA\Desktop\ChamaBot\quotes.docx")

# ✅ Save the extracted Q&A to a JSON file
with open("chama_data.json", "w") as f:
    json.dump(qa_data, f, indent=4)

# Save quotes
with open("quotes.json", "w") as f:
    json.dump(quotes, f, indent=4)

print("✅ Data extraction complete!")
